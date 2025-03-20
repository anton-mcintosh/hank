import os
import uuid
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

load_dotenv()

# Configuration from environment variables
INVOICE_DIR = os.getenv("INVOICE_DIR", "./invoices")
COMPANY_NAME = os.getenv("COMPANY_NAME", "Auto Shop")
COMPANY_ADDRESS = os.getenv("COMPANY_ADDRESS", "123 Main St, Anytown, USA")
COMPANY_PHONE = os.getenv("COMPANY_PHONE", "(555) 123-4567")
COMPANY_EMAIL = os.getenv("COMPANY_EMAIL", "service@autoshop.com")
COMPANY_WEBSITE = os.getenv("COMPANY_WEBSITE", "www.yourautoshop.com")
LOGO_PATH = os.getenv("LOGO_PATH", None)  # Optional path to logo

# Ensure invoice directory exists
os.makedirs(INVOICE_DIR, exist_ok=True)

# Helper function to set cell background color
def set_cell_background(cell, color):
    """Set the background color of a table cell"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

async def generate_invoice_docx(work_order, customer=None, vehicle=None, is_estimate=False):
    """Generate an editable Word document invoice
    
    Args:
        work_order: WorkOrder database object
        customer: Customer database object (optional)
        vehicle: Vehicle database object (optional)
        is_estimate: Boolean indicating if this is an estimate (True) or invoice (False)
        
    Returns:
        str: Path to the generated DOCX file
    """
    try:
        # Set up file path and document type
        document_type = "Estimate" if is_estimate else "Invoice"
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        filename = f"{document_type}_{work_order.id[:8]}_{timestamp}.docx"
        file_path = os.path.join(INVOICE_DIR, filename)
        
        # Create a new Document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"{document_type} #{work_order.id[:8]}"
        doc.core_properties.author = COMPANY_NAME
        
        # ----- HEADER SECTION -----
        # Create header table for company info and logo
        header_table = doc.add_table(rows=1, cols=2)
        header_table.style = 'Table Grid'
        header_table.autofit = False
        header_table.allow_autofit = False
        
        # Set column widths
        for cell in header_table.columns[0].cells:
            cell.width = Inches(4.0)
        for cell in header_table.columns[1].cells:
            cell.width = Inches(2.5)
            
        # Left cell: Company info
        company_cell = header_table.cell(0, 0)
        company_info = company_cell.paragraphs[0]
        company_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add company name with bold formatting
        company_name_run = company_info.add_run(COMPANY_NAME)
        company_name_run.bold = True
        company_name_run.font.size = Pt(14)
        company_info.add_run("\n" + COMPANY_ADDRESS)
        company_info.add_run("\nPhone: " + COMPANY_PHONE)
        company_info.add_run("\nEmail: " + COMPANY_EMAIL)
        if COMPANY_WEBSITE:
            company_info.add_run("\nWeb: " + COMPANY_WEBSITE)
            
        # Right cell: Document info
        doc_cell = header_table.cell(0, 1)
        doc_info = doc_cell.paragraphs[0]
        doc_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add document title with bold formatting
        doc_title_run = doc_info.add_run(f"{document_type.upper()} #{work_order.id[:8]}")
        doc_title_run.bold = True
        doc_title_run.font.size = Pt(14)
        doc_info.add_run(f"\nDate: {datetime.now().strftime('%m/%d/%Y')}")
        doc_info.add_run(f"\nStatus: {work_order.status.upper()}")
        
        # Add some space after header
        doc.add_paragraph()
        
        # ----- CUSTOMER AND VEHICLE INFO -----
        # Customer information section
        customer_heading = doc.add_heading("Customer Information", level=2)
        
        # Create a table for customer info
        customer_table = doc.add_table(rows=4, cols=2)
        customer_table.style = 'Table Grid'
        
        # Set column widths
        for cell in customer_table.columns[0].cells:
            cell.width = Inches(1.5)
        
        # Add customer data
        rows = customer_table.rows
        
        # Set header row
        header_row = rows[0]
        header_row.cells[0].text = "Customer:"
        header_row.cells[0].paragraphs[0].runs[0].bold = True
        
        if customer:
            header_row.cells[1].text = f"{customer.first_name} {customer.last_name}"
            rows[1].cells[0].text = "Phone:"
            rows[1].cells[0].paragraphs[0].runs[0].bold = True
            rows[1].cells[1].text = customer.phone
            
            rows[2].cells[0].text = "Email:"
            rows[2].cells[0].paragraphs[0].runs[0].bold = True
            rows[2].cells[1].text = customer.email
            
            rows[3].cells[0].text = "Address:"
            rows[3].cells[0].paragraphs[0].runs[0].bold = True
            rows[3].cells[1].text = customer.address if customer.address else ""
        else:
            header_row.cells[1].text = work_order.customer_name or "Unknown"
            rows[1].cells[0].text = "Phone:"
            rows[1].cells[0].paragraphs[0].runs[0].bold = True
            rows[1].cells[1].text = "N/A"
            
            rows[2].cells[0].text = "Email:"
            rows[2].cells[0].paragraphs[0].runs[0].bold = True
            rows[2].cells[1].text = "N/A"
            
            rows[3].cells[0].text = "Address:"
            rows[3].cells[0].paragraphs[0].runs[0].bold = True
            rows[3].cells[1].text = "N/A"
        
        # Add some space
        doc.add_paragraph()
        
        # Vehicle information section
        vehicle_heading = doc.add_heading("Vehicle Information", level=2)
        
        # Create a table for vehicle info
        vehicle_table = doc.add_table(rows=3, cols=2)
        vehicle_table.style = 'Table Grid'
        
        # Set column widths
        for cell in vehicle_table.columns[0].cells:
            cell.width = Inches(1.5)
        
        # Add vehicle data
        rows = vehicle_table.rows
        
        # Prepare vehicle info
        if vehicle:
            vehicle_year = vehicle.year if vehicle.year else "Unknown"
            vehicle_make = vehicle.make if vehicle.make else "Unknown"
            vehicle_model = vehicle.model if vehicle.model else "Unknown"
            
            rows[0].cells[0].text = "Year/Make/Model:"
            rows[0].cells[0].paragraphs[0].runs[0].bold = True
            rows[0].cells[1].text = f"{vehicle_year} {vehicle_make} {vehicle_model}"
            
            rows[1].cells[0].text = "VIN:"
            rows[1].cells[0].paragraphs[0].runs[0].bold = True
            rows[1].cells[1].text = vehicle.vin if vehicle.vin else "Unknown"
            
            rows[2].cells[0].text = "Mileage:"
            rows[2].cells[0].paragraphs[0].runs[0].bold = True
            rows[2].cells[1].text = f"{vehicle.mileage:,}" if vehicle.mileage else "Unknown"
        else:
            # Use vehicle_info from work_order if available
            vi = work_order.vehicle_info or {}
            
            rows[0].cells[0].text = "Year/Make/Model:"
            rows[0].cells[0].paragraphs[0].runs[0].bold = True
            rows[0].cells[1].text = f"{vi.get('year', 'Unknown')} {vi.get('make', 'Unknown')} {vi.get('model', 'Unknown')}"
            
            rows[1].cells[0].text = "VIN:"
            rows[1].cells[0].paragraphs[0].runs[0].bold = True
            rows[1].cells[1].text = vi.get('vin', 'Unknown')
            
            rows[2].cells[0].text = "Mileage:"
            rows[2].cells[0].paragraphs[0].runs[0].bold = True
            rows[2].cells[1].text = f"{vi.get('mileage', 'Unknown'):,}" if vi.get('mileage') else "Unknown"
        
        # Add some space
        doc.add_paragraph()
        
        # ----- WORK SUMMARY -----
        doc.add_heading("Work Summary", level=2)
        doc.add_paragraph(work_order.work_summary)
        
        # Add some space
        doc.add_paragraph()
        
        # ----- LINE ITEMS TABLE -----
        doc.add_heading("Services & Parts", level=2)
        
        # Create table for line items
        line_items = work_order.line_items or []
        # Add 4 rows: header + items + totals (parts, labor, grand total)
        table_rows = len(line_items) + 4
        line_items_table = doc.add_table(rows=table_rows, cols=5)
        line_items_table.style = 'Table Grid'
        line_items_table.autofit = False
        
        # Set column widths
        col_widths = [3.5, 1.0, 1.0, 1.0, 1.0]  # in inches
        for i, width in enumerate(col_widths):
            for cell in line_items_table.columns[i].cells:
                cell.width = Inches(width)
        
        # Header row
        header_row = line_items_table.rows[0]
        headers = ["Description", "Type", "Quantity", "Price", "Total"]
        
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
            # Dark blue background for header
            set_cell_background(cell, "3B71CA")  # Dark blue
            # White text for header
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add line items
        for i, item in enumerate(line_items):
            row = line_items_table.rows[i + 1]
            
            # Description
            row.cells[0].text = item.get('description', '')
            
            # Type
            row.cells[1].text = item.get('type', '').capitalize()
            
            # Quantity
            quantity = f"{item.get('quantity', 0):.1f}" if isinstance(item.get('quantity'), (int, float)) else str(item.get('quantity', ''))
            row.cells[2].text = quantity
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Unit Price
            unit_price = f"${item.get('unit_price', 0):.2f}" if isinstance(item.get('unit_price'), (int, float)) else str(item.get('unit_price', ''))
            row.cells[3].text = unit_price
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Total
            total = f"${item.get('total', 0):.2f}" if isinstance(item.get('total'), (int, float)) else str(item.get('total', ''))
            row.cells[4].text = total
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Zebra striping for rows
            if i % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, "F5F5F5")  # Light gray
        
        # Calculate the row index for totals
        totals_start_idx = len(line_items) + 1
        
        # Parts Total
        parts_row = line_items_table.rows[totals_start_idx]
        parts_row.cells[0].text = ""
        parts_row.cells[1].text = ""
        parts_row.cells[2].text = ""
        parts_row.cells[3].text = "Parts Total:"
        parts_row.cells[3].paragraphs[0].runs[0].bold = True
        parts_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        parts_total = f"${work_order.total_parts:.2f}" if isinstance(work_order.total_parts, (int, float)) else str(work_order.total_parts)
        parts_row.cells[4].text = parts_total
        parts_row.cells[4].paragraphs[0].runs[0].bold = True
        parts_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Labor Total
        labor_row = line_items_table.rows[totals_start_idx + 1]
        labor_row.cells[0].text = ""
        labor_row.cells[1].text = ""
        labor_row.cells[2].text = ""
        labor_row.cells[3].text = "Labor Total:"
        labor_row.cells[3].paragraphs[0].runs[0].bold = True
        labor_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        labor_total = f"${work_order.total_labor:.2f}" if isinstance(work_order.total_labor, (int, float)) else str(work_order.total_labor)
        labor_row.cells[4].text = labor_total
        labor_row.cells[4].paragraphs[0].runs[0].bold = True
        labor_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Grand Total
        grand_row = line_items_table.rows[totals_start_idx + 2]
        grand_row.cells[0].text = ""
        grand_row.cells[1].text = ""
        grand_row.cells[2].text = ""
        grand_row.cells[3].text = "GRAND TOTAL:"
        grand_row.cells[3].paragraphs[0].runs[0].bold = True
        grand_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        grand_total = f"${work_order.total:.2f}" if isinstance(work_order.total, (int, float)) else str(work_order.total)
        grand_row.cells[4].text = grand_total
        grand_row.cells[4].paragraphs[0].runs[0].bold = True
        grand_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add some space
        doc.add_paragraph()
        
        # ----- NOTES SECTION -----
        if is_estimate:
            note_para = doc.add_paragraph()
            note_run = note_para.add_run("PLEASE NOTE: This is an ESTIMATE only. Actual charges may vary based on additional parts or labor required. This estimate is valid for 30 days.")
            note_run.bold = True
        else:
            # Payment terms
            terms_para = doc.add_paragraph()
            terms_run = terms_para.add_run("PAYMENT TERMS: Payment due upon completion of service. We accept cash, checks, and all major credit cards.")
            terms_run.bold = True
        
        # Thank you message
        thanks_para = doc.add_paragraph()
        thanks_para.add_run("Thank you for your business!")
        
        # ----- FOOTER -----
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(f"{COMPANY_NAME} • {COMPANY_PHONE} • {COMPANY_EMAIL}")
        
        # Save the document
        doc.save(file_path)
        return file_path
        
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None
